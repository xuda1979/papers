# -*- coding: utf-8 -*-
"""
提取docx中的图片并分析它们在文档中的位置
"""
from docx import Document
from docx.oxml.ns import qn
import os
import zipfile
import shutil

def extract_images_from_docx(docx_path, output_dir):
    """从docx文件中提取所有图片"""
    
    # 创建输出目录
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    # docx本质上是一个zip文件
    with zipfile.ZipFile(docx_path, 'r') as zip_ref:
        # 获取所有媒体文件
        media_files = [f for f in zip_ref.namelist() if f.startswith('word/media/')]
        
        print(f"找到 {len(media_files)} 个媒体文件")
        
        for media_file in media_files:
            # 提取文件名
            filename = os.path.basename(media_file)
            # 转换gif为png以便LaTeX使用
            output_filename = filename
            
            # 提取文件
            with zip_ref.open(media_file) as src:
                output_path = os.path.join(output_dir, output_filename)
                with open(output_path, 'wb') as dst:
                    dst.write(src.read())
            print(f"已提取: {output_filename}")
    
    return len(media_files)

def get_image_positions_in_doc(docx_path):
    """分析图片在文档中的位置（段落顺序）"""
    doc = Document(docx_path)
    
    image_info = []
    
    for i, para in enumerate(doc.paragraphs):
        # 获取段落文本
        text = para.text.strip()
        
        # 检查段落中是否有图片
        for run in para.runs:
            # 查找drawing元素（图片）
            drawings = run._element.findall('.//' + qn('w:drawing'))
            for drawing in drawings:
                # 尝试获取图片引用
                blips = drawing.findall('.//' + qn('a:blip'))
                for blip in blips:
                    embed_id = blip.get(qn('r:embed'))
                    if embed_id:
                        # 获取图片文件名
                        rel = doc.part.rels.get(embed_id)
                        if rel:
                            image_name = os.path.basename(rel.target_ref)
                            # 获取上下文（前后段落）
                            prev_text = doc.paragraphs[i-1].text.strip() if i > 0 else ""
                            next_text = doc.paragraphs[i+1].text.strip() if i < len(doc.paragraphs)-1 else ""
                            
                            image_info.append({
                                'paragraph_index': i,
                                'image_name': image_name,
                                'current_text': text[:100] if text else "(无文本)",
                                'prev_text': prev_text[:80] if prev_text else "",
                                'next_text': next_text[:80] if next_text else ""
                            })
    
    return image_info

if __name__ == '__main__':
    docx_path = 'QUBO 解决通信痛点研究.docx'
    output_dir = 'figures'
    
    # 提取图片
    count = extract_images_from_docx(docx_path, output_dir)
    print(f"\n总共提取了 {count} 张图片到 {output_dir} 目录")
    
    # 分析图片位置
    print("\n\n=== 图片在文档中的位置 ===\n")
    positions = get_image_positions_in_doc(docx_path)
    
    for info in positions:
        print(f"图片: {info['image_name']}")
        print(f"  段落索引: {info['paragraph_index']}")
        if info['prev_text']:
            print(f"  前文: {info['prev_text'][:60]}...")
        if info['current_text'] != "(无文本)":
            print(f"  当前: {info['current_text'][:60]}...")
        if info['next_text']:
            print(f"  后文: {info['next_text'][:60]}...")
        print()
