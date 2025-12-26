# -*- coding: utf-8 -*-
"""
详细分析docx中主要插图（PNG）的位置和上下文
"""
from docx import Document
from docx.oxml.ns import qn
import os

def analyze_png_images(docx_path):
    """分析PNG图片在文档中的位置"""
    doc = Document(docx_path)
    
    # 先打印所有段落内容和索引
    print("=== 文档段落概览 ===\n")
    
    # 收集图片信息
    image_contexts = []
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        # 检查段落中是否有图片
        for run in para.runs:
            drawings = run._element.findall('.//' + qn('w:drawing'))
            for drawing in drawings:
                blips = drawing.findall('.//' + qn('a:blip'))
                for blip in blips:
                    embed_id = blip.get(qn('r:embed'))
                    if embed_id:
                        rel = doc.part.rels.get(embed_id)
                        if rel:
                            image_name = os.path.basename(rel.target_ref)
                            # 只关注PNG图片（实际插图）
                            if image_name.endswith('.png'):
                                # 获取周围段落的上下文
                                context_before = []
                                context_after = []
                                
                                # 前5个段落
                                for j in range(max(0, i-5), i):
                                    t = doc.paragraphs[j].text.strip()
                                    if t:
                                        context_before.append(f"[{j}] {t[:100]}")
                                
                                # 后5个段落
                                for j in range(i+1, min(len(doc.paragraphs), i+6)):
                                    t = doc.paragraphs[j].text.strip()
                                    if t:
                                        context_after.append(f"[{j}] {t[:100]}")
                                
                                image_contexts.append({
                                    'image': image_name,
                                    'para_idx': i,
                                    'before': context_before,
                                    'current': text[:150] if text else "(无文本)",
                                    'after': context_after
                                })
    
    print(f"找到 {len(image_contexts)} 张PNG插图\n")
    print("=" * 80)
    
    for ctx in image_contexts:
        print(f"\n【图片: {ctx['image']}】 (段落 {ctx['para_idx']})")
        print("-" * 40)
        print("前文:")
        for line in ctx['before']:
            print(f"  {line}")
        print(f"\n当前段落: {ctx['current']}")
        print("\n后文:")
        for line in ctx['after']:
            print(f"  {line}")
        print("=" * 80)

if __name__ == '__main__':
    analyze_png_images('QUBO 解决通信痛点研究.docx')
